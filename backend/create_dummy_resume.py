import docx

doc = docx.Document()
doc.add_heading('John Doe - Software Engineer Resume', 0)

doc.add_heading('Education', level=1)
doc.add_paragraph('Bachelor of Science in Computer Science\nXYZ University, 2020 - 2024')

doc.add_heading('Work Experience', level=1)
p = doc.add_paragraph()
p.add_run('Software Engineer Intern\nABC Solutions, Jan 2023 - Present\n')
p.add_run('- Developed web applications using Python and Flask.\n')
p.add_run('- Created responsive user interfaces with HTML, CSS, and basic JavaScript.\n')
p.add_run('- Wrote SQL queries and worked with databases.')

doc.add_heading('Projects', level=1)
p2 = doc.add_paragraph()
p2.add_run('Personal Portfolio Web App\n')
p2.add_run('- Built a personal portfolio website in HTML and CSS.\n')
p2.add_run('- Hosted the site on GitHub Pages.')

doc.add_heading('Skills', level=1)
doc.add_paragraph('Python, Flask, JavaScript, HTML, CSS, SQLite, Git, SQL')

doc.save('dummy_resume.docx')
print("dummy_resume.docx created successfully!")
